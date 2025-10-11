# Export utilities for formatted documents
# Create professionally formatted DOCX and PDF files

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from io import BytesIO
import re

def create_docx(formatted_letter):
    """
    Create a professionally formatted Word document
    Returns BytesIO buffer ready for download
    """
    print("📝 Creating DOCX file...")
    
    doc = Document()
    
    # Set margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse the letter
    lines = [l.strip() for l in formatted_letter.split('\n') if l.strip()]
    
    # Extract header section (name, address, contact, date)
    header_lines = []
    body_start = 0
    
    for i, line in enumerate(lines):
        if line.startswith('Dear'):
            body_start = i
            break
        header_lines.append(line)
    
    # Add name (first line) - Centered, Bold, Large
    if header_lines:
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(header_lines[0])
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(6)
    
    # Add address and contact info - Centered, Regular
    for line in header_lines[1:]:
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(line)
        contact_run.font.name = 'Calibri'
        contact_run.font.size = Pt(11)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(3)
    
    # Add spacing before date
    doc.add_paragraph()
    
    # Add body paragraphs - Justified
    current_paragraph = []
    for i in range(body_start, len(lines)):
        line = lines[i]
        
        # Check if this is a paragraph break (salutation, closing, etc.)
        if line.startswith('Dear') or line in ['Sincerely,', 'Best regards,', 'Regards,']:
            # Add accumulated paragraph first
            if current_paragraph:
                para_text = ' '.join(current_paragraph)
                para = doc.add_paragraph(para_text)
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(12)
                current_paragraph = []
            
            # Add the special line (Dear, Sincerely, etc.)
            para = doc.add_paragraph(line)
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(12)
        else:
            # Accumulate paragraph text
            current_paragraph.append(line)
    
    # Add any remaining paragraph
    if current_paragraph:
        para_text = ' '.join(current_paragraph)
        para = doc.add_paragraph(para_text)
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(12)
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    print("✓ DOCX created")
    
    return buffer


def create_pdf(formatted_letter):
    """
    Create a professionally formatted PDF document
    Returns bytes ready for download
    """
    print("📕 Creating PDF file...")
    
    pdf = FPDF()
    pdf.add_page()
    
    # Set margins
    pdf.set_margins(left=25, top=25, right=25)
    
    # Parse letter
    lines = [l.strip() for l in formatted_letter.split('\n') if l.strip()]
    
    # Extract header section
    header_lines = []
    body_start = 0
    
    for i, line in enumerate(lines):
        if line.startswith('Dear'):
            body_start = i
            break
        header_lines.append(line)
    
    # Add name - Centered, Bold, Large
    if header_lines:
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, header_lines[0], ln=True, align='C')
        pdf.ln(2)
    
    # Add contact info - Centered, Regular
    pdf.set_font('Arial', '', 11)
    for line in header_lines[1:]:
        pdf.cell(0, 6, line, ln=True, align='C')
    
    pdf.ln(8)  # Space before date
    
    # Process body
    pdf.set_font('Arial', '', 11)
    
    current_paragraph = []
    
    for i in range(body_start, len(lines)):
        line = lines[i]
        
        # Check if this is a paragraph break
        if line.startswith('Dear') or line in ['Sincerely,', 'Best regards,', 'Regards,', 'Thank you,']:
            # Output accumulated paragraph first
            if current_paragraph:
                para_text = ' '.join(current_paragraph)
                pdf.multi_cell(0, 6, para_text, align='J')  # Justified
                pdf.ln(4)
                current_paragraph = []
            
            # Add the special line
            pdf.cell(0, 6, line, ln=True, align='L')
            pdf.ln(4)
        else:
            # Accumulate paragraph
            current_paragraph.append(line)
    
    # Add any remaining paragraph
    if current_paragraph:
        para_text = ' '.join(current_paragraph)
        pdf.multi_cell(0, 6, para_text, align='J')  # Justified
    
    # Return PDF as bytes - FIXED: Use dest='S' to return string/bytes
    pdf_output = pdf.output(dest='S')
    
    # Convert to bytes if it's a string
    if isinstance(pdf_output, str):
        pdf_bytes = pdf_output.encode('latin-1')
    else:
        pdf_bytes = pdf_output
    
    print("✓ PDF created (bytes returned)")
    
    return pdf_bytes